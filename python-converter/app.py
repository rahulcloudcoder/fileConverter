from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile
from werkzeug.utils import secure_filename
import traceback
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pythoncom
import win32com.client
import subprocess
import time
from PIL import Image
import io
import re

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "service": "Professional Document Converter",
        "capabilities": ["pdf-to-docx", "docx-to-pdf", "txt-conversion"]
    })

@app.route('/convert', methods=['POST'])
def convert_file():
    input_path = None
    output_path = None
    
    try:
        print("=== Starting Professional Conversion ===")
        
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        target_format = request.form.get('targetFormat', '').lower()
        
        print(f"File: {file.filename}")
        print(f"Target format: {target_format}")
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"error": f"File type not allowed. Allowed: {ALLOWED_EXTENSIONS}"}), 400
        
        if not target_format:
            return jsonify({"error": "Target format not specified"}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"input_{int(time.time())}_{filename}")
        file.save(input_path)
        
        print(f"Input saved: {input_path}")
        print(f"File size: {os.path.getsize(input_path)} bytes")
        
        # Generate output path
        base_name = os.path.splitext(filename)[0]
        output_filename = f"converted_{base_name}.{target_format}"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"output_{int(time.time())}_{output_filename}")
        
        # Perform conversion
        result_path = perform_conversion(input_path, output_path, target_format)
        
        if result_path and os.path.exists(result_path):
            file_size = os.path.getsize(result_path)
            print(f"Conversion SUCCESS! Output: {result_path} ({file_size} bytes)")
            
            return send_file(
                result_path,
                as_attachment=True,
                download_name=output_filename,
                mimetype=get_mimetype(target_format)
            )
        else:
            return jsonify({"error": "Conversion failed - no output file created"}), 500
            
    except Exception as e:
        print(f"❌ Conversion ERROR: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500
    finally:
        # Cleanup files
        cleanup_file(input_path)
        cleanup_file(output_path)

def perform_conversion(input_path, output_path, target_format):
    """Main conversion router"""
    input_ext = os.path.splitext(input_path)[1].lower()
    
    print(f"Converting {input_ext} to {target_format}")
    
    # PDF to DOCX - Use professional converter
    if input_ext == '.pdf' and target_format == 'docx':
        return pdf_to_docx_professional(input_path, output_path)
    
    # PDF to TXT
    elif input_ext == '.pdf' and target_format == 'txt':
        return pdf_to_txt_advanced(input_path, output_path)
    
    # DOCX to PDF
    elif input_ext == '.docx' and target_format == 'pdf':
        return docx_to_pdf_professional(input_path, output_path)
    
    # DOCX to TXT
    elif input_ext == '.docx' and target_format == 'txt':
        return docx_to_txt_advanced(input_path, output_path)
    
    else:
        raise ValueError(f"Unsupported conversion: {input_ext} to {target_format}")

def pdf_to_docx_professional(pdf_path, docx_path):
    """Professional PDF to DOCX conversion with images and formatting"""
    try:
        print("🔄 Starting professional PDF to DOCX conversion...")
        
        # Method 1: Try pdf2docx first (best for layout preservation)
        try:
            from pdf2docx import Converter
            print("Trying pdf2docx converter...")
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            
            # Verify conversion worked
            if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
                print("✅ pdf2docx conversion successful!")
                return docx_path
            else:
                print("❌ pdf2docx produced empty file, trying fallback...")
        except Exception as e:
            print(f"❌ pdf2docx failed: {e}, trying advanced fallback...")
        
        # Method 2: Advanced fallback with PyMuPDF
        return pdf_to_docx_advanced_fallback(pdf_path, docx_path)
        
    except Exception as e:
        raise Exception(f"Professional PDF to DOCX conversion failed: {str(e)}")

def pdf_to_docx_advanced_fallback(pdf_path, docx_path):
    """Advanced fallback with better formatting and image extraction"""
    try:
        print("Using advanced PyMuPDF fallback...")
        
        pdf_doc = fitz.open(pdf_path)
        word_doc = Document()
        
        # Set document properties
        word_doc.core_properties.title = "Converted PDF Document"
        
        # Process each page
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            print(f"Processing page {page_num + 1}...")
            
            # Add page header
            if page_num > 0:
                word_doc.add_page_break()
            
            # Add page number indicator
            page_header = word_doc.add_paragraph()
            page_header_run = page_header.add_run(f"--- Page {page_num + 1} ---")
            page_header_run.font.size = Pt(9)
            page_header_run.font.color.rgb = RGBColor(128, 128, 128)
            page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Extract images from the page
            image_list = page.get_images()
            image_count = 0
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # RGB or CMYK
                        img_data = pix.tobytes("png")
                        
                        # Save image temporarily
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_img:
                            temp_img.write(img_data)
                            temp_img_path = temp_img.name
                        
                        # Add image to document
                        paragraph = word_doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        
                        # Add image with reasonable size
                        run.add_picture(temp_img_path, width=Inches(5.0))
                        
                        # Add image caption
                        caption = word_doc.add_paragraph()
                        caption_run = caption.add_run(f"[Image {image_count + 1}]")
                        caption_run.font.size = Pt(9)
                        caption_run.font.italic = True
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        image_count += 1
                        
                        # Cleanup temp file
                        os.unlink(temp_img_path)
                    
                    pix = None
                    
                except Exception as img_error:
                    print(f"Image extraction error: {img_error}")
                    continue
            
            # Extract text with formatting information
            text_blocks = page.get_text("dict")
            
            for block in text_blocks.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        paragraph = word_doc.add_paragraph()
                        
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if not text:
                                continue
                            
                            run = paragraph.add_run(text + " ")
                            
                            # Apply font formatting
                            font_size = span["size"]
                            if font_size > 12:
                                run.font.size = Pt(min(font_size, 36))  # Cap at 36pt
                            
                            # Bold detection
                            if span["flags"] & 2:  # Bold flag
                                run.font.bold = True
                            
                            # Italic detection
                            if span["flags"] & 1:  # Italic flag
                                run.font.italic = True
                            
                            # Font color (basic conversion)
                            if span["color"] != 0:
                                try:
                                    # Convert color (simplified)
                                    color = span["color"]
                                    if color != 0:
                                        run.font.color.rgb = RGBColor(0, 0, 0)  # Default to black
                                except:
                                    pass
                        
                        # Set paragraph formatting
                        if paragraph.text.strip():
                            paragraph.paragraph_format.space_after = Pt(8)
                            paragraph.paragraph_format.line_spacing = 1.15
            
            print(f"Page {page_num + 1} completed - {image_count} images extracted")
        
        # Save the document
        word_doc.save(docx_path)
        pdf_doc.close()
        
        print("✅ Advanced PDF to DOCX conversion completed successfully")
        return docx_path
        
    except Exception as e:
        raise Exception(f"Advanced PDF to DOCX fallback failed: {str(e)}")

def pdf_to_txt_advanced(pdf_path, txt_path):
    """Advanced PDF to text conversion with layout preservation"""
    try:
        print("Converting PDF to TXT (advanced)...")
        
        pdf_doc = fitz.open(pdf_path)
        full_text = []
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            # Get text with layout preservation
            text = page.get_text("text", sort=True)
            
            if text.strip():
                full_text.append(f"\n{'='*50}")
                full_text.append(f"PAGE {page_num + 1}")
                full_text.append(f"{'='*50}\n")
                full_text.append(text.strip())
        
        pdf_doc.close()
        
        # Clean and format text
        cleaned_text = clean_extracted_text('\n'.join(full_text))
        
        # Write to file
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(cleaned_text)
        
        print("PDF to TXT conversion completed successfully")
        return txt_path
        
    except Exception as e:
        raise Exception(f"PDF to TXT conversion failed: {str(e)}")

def docx_to_txt_advanced(docx_path, txt_path):
    """Advanced DOCX to text conversion with formatting awareness"""
    try:
        print("Converting DOCX to TXT (advanced)...")
        
        doc = Document(docx_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Detect headings by style
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.split()[-1] if 'Heading' in paragraph.style.name else '1'
                    full_text.append(f"\n{'#' * int(level)} {text}")
                else:
                    full_text.append(text)
        
        # Extract tables
        for table in doc.tables:
            full_text.append("\n" + "-" * 40)
            full_text.append("TABLE:")
            full_text.append("-" * 40)
            
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
            
            full_text.append("-" * 40)
        
        # Write to file
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        
        print("DOCX to TXT conversion completed successfully")
        return txt_path
        
    except Exception as e:
        raise Exception(f"DOCX to TXT conversion failed: {str(e)}")

def docx_to_pdf_professional(docx_path, pdf_path):
    """Professional DOCX to PDF conversion"""
    try:
        print("Converting DOCX to PDF...")
        
        # Method 1: Try Windows Word first (best quality)
        if os.name == 'nt':
            try:
                return docx_to_pdf_windows(docx_path, pdf_path)
            except Exception as e:
                print(f"Windows Word conversion failed: {e}")
        
        # Method 2: Try LibreOffice
        try:
            return docx_to_pdf_libreoffice(docx_path, pdf_path)
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")
        
        # Method 3: Use reportlab as fallback
        return docx_to_pdf_reportlab(docx_path, pdf_path)
        
    except Exception as e:
        raise Exception(f"DOCX to PDF conversion failed: {str(e)}")

def docx_to_pdf_windows(docx_path, pdf_path):
    """Use Microsoft Word on Windows (best quality)"""
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Use absolute paths
        abs_docx_path = os.path.abspath(docx_path)
        abs_pdf_path = os.path.abspath(pdf_path)
        
        doc = word.Documents.Open(abs_docx_path)
        doc.SaveAs(abs_pdf_path, FileFormat=17)  # 17 = PDF format
        doc.Close()
        word.Quit()
        
        print("DOCX to PDF conversion completed (Windows Word - Professional)")
        return pdf_path
    except Exception as e:
        raise Exception(f"Windows Word conversion error: {str(e)}")
    finally:
        try:
            pythoncom.CoUninitialize()
        except:
            pass

def docx_to_pdf_libreoffice(docx_path, pdf_path):
    """Use LibreOffice for conversion"""
    try:
        output_dir = os.path.dirname(pdf_path)
        
        # Try different LibreOffice command variations
        commands = [
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            ['/usr/bin/libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path]
        ]
        
        for cmd in commands:
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                if result.returncode == 0:
                    # Find the generated PDF
                    base_name = os.path.splitext(os.path.basename(docx_path))[0]
                    expected_pdf = os.path.join(output_dir, f"{base_name}.pdf")
                    
                    if os.path.exists(expected_pdf):
                        os.rename(expected_pdf, pdf_path)
                        print("DOCX to PDF conversion completed (LibreOffice)")
                        return pdf_path
            except (subprocess.TimeoutExpired, FileNotFoundError):
                continue
        
        raise Exception("LibreOffice not found or conversion failed")
        
    except Exception as e:
        raise Exception(f"LibreOffice conversion error: {str(e)}")

def docx_to_pdf_reportlab(docx_path, pdf_path):
    """Fallback using reportlab (basic but reliable)"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import ImageReader
        
        print("Using ReportLab fallback for DOCX to PDF...")
        
        doc = Document(docx_path)
        c = canvas.Canvas(pdf_path, pagesize=letter)
        
        width, height = letter
        y_position = height - 50
        line_height = 14
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Simple text wrapping
                words = text.split()
                lines = []
                current_line = []
                
                for word in words:
                    current_line.append(word)
                    test_line = ' '.join(current_line)
                    if len(test_line) > 80:
                        lines.append(' '.join(current_line[:-1]))
                        current_line = [word]
                
                if current_line:
                    lines.append(' '.join(current_line))
                
                # Add lines to PDF
                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    
                    c.drawString(50, y_position, line)
                    y_position -= line_height
                
                # Space between paragraphs
                y_position -= line_height / 2
        
        c.save()
        print("DOCX to PDF conversion completed (ReportLab Fallback)")
        return pdf_path
        
    except Exception as e:
        raise Exception(f"ReportLab conversion error: {str(e)}")

def clean_extracted_text(text):
    """Clean and format extracted text"""
    if not text:
        return "No text content found."
    
    # Remove excessive whitespace but preserve paragraph breaks
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            cleaned_lines.append(line)
        elif cleaned_lines and cleaned_lines[-1] != '':
            cleaned_lines.append('')
    
    # Remove multiple consecutive empty lines
    result = []
    for i, line in enumerate(cleaned_lines):
        if line == '' and result and result[-1] == '':
            continue
        result.append(line)
    
    return '\n'.join(result)

def get_mimetype(format_name):
    """Get MIME type for file format"""
    mimetypes = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain',
        'doc': 'application/msword',
    }
    return mimetypes.get(format_name, 'application/octet-stream')

def cleanup_file(file_path):
    """Cleanup a single file with retry"""
    if not file_path or not os.path.exists(file_path):
        return
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            os.remove(file_path)
            print(f"Cleaned up: {file_path}")
            break
        except Exception as e:
            if attempt == max_retries - 1:
                print(f"Warning: Could not delete {file_path}: {e}")
            time.sleep(0.1)

if __name__ == '__main__':
    print("🚀 Starting Professional Document Converter Service...")
    print(f"📁 Upload folder: {UPLOAD_FOLDER}")
    print("✅ Service ready! Waiting for conversion requests...")
    app.run(host='0.0.0.0', port=8000, debug=True)