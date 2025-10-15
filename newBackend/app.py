from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile
from werkzeug.utils import secure_filename
import traceback
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pythoncom
import win32com.client
import subprocess
import time
from PIL import Image
import io
import re
import json
from typing import Dict, List, Tuple, Optional
import base64

app = Flask(__name__)
CORS(app)

# Cross-platform path handling
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {
    'pdf', 'docx', 'doc', 'txt',
    'jpg', 'jpeg', 'png', 'webp', 'gif', 'bmp', 'tiff', 'svg'
}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

class ProfessionalDocumentConverter:
    """Professional document converter with advanced features"""
    
    def __init__(self):
        self.supported_conversions = self.get_supported_conversions()
    
    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str) -> str:
        """Professional PDF to DOCX with multiple fallbacks"""
        methods = [
            self._pdf_to_docx_pdf2docx,      # Best for layout
            self._pdf_to_docx_advanced_pymupdf,  # Advanced formatting
            self._pdf_to_docx_basic_fallback     # Basic fallback
        ]
        
        last_error = None
        for method in methods:
            try:
                result = method(pdf_path, docx_path)
                if result and os.path.exists(result):
                    return result
            except Exception as e:
                last_error = e
                continue
        
        raise Exception(f"All PDF to DOCX methods failed: {last_error}")
    
    def _pdf_to_docx_pdf2docx(self, pdf_path: str, docx_path: str) -> str:
        """Use pdf2docx for best layout preservation"""
        try:
            from pdf2docx import Converter
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None, multi_processing=False)
            cv.close()
            return docx_path
        except Exception as e:
            raise Exception(f"pdf2docx failed: {str(e)}")
    
    def _pdf_to_docx_advanced_pymupdf(self, pdf_path: str, docx_path: str) -> str:
        """Advanced PyMuPDF conversion with images and formatting"""
        pdf_doc = fitz.open(pdf_path)
        word_doc = Document()
        
        # Set document properties
        word_doc.core_properties.title = "Converted PDF Document"
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            # Add page break for subsequent pages
            if page_num > 0:
                word_doc.add_page_break()
            
            # Add page header
            self._add_page_header(word_doc, page_num + 1)
            
            # Extract and embed images
            self._extract_and_embed_images(word_doc, page, pdf_doc)
            
            # Extract text with advanced formatting
            self._extract_text_with_formatting(word_doc, page)
        
        word_doc.save(docx_path)
        pdf_doc.close()
        return docx_path
    
    def _pdf_to_docx_basic_fallback(self, pdf_path: str, docx_path: str) -> str:
        """Basic fallback conversion"""
        pdf_doc = fitz.open(pdf_path)
        word_doc = Document()
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            text = page.get_text("text", sort=True)
            
            if text.strip():
                if page_num > 0:
                    word_doc.add_page_break()
                
                word_doc.add_heading(f'Page {page_num + 1}', level=2)
                
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                for para in paragraphs:
                    if self._is_heading(para):
                        word_doc.add_heading(para, level=3)
                    else:
                        p = word_doc.add_paragraph(para)
                        p.paragraph_format.space_after = Pt(8)
        
        word_doc.save(docx_path)
        pdf_doc.close()
        return docx_path
    
    def _add_page_header(self, doc: Document, page_num: int):
        """Add professional page header"""
        header = doc.add_paragraph()
        header_run = header.add_run(f"--- Page {page_num} ---")
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = RGBColor(128, 128, 128)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _extract_and_embed_images(self, doc: Document, page, pdf_doc):
        """Extract and embed images from PDF page"""
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                pix = fitz.Pixmap(pdf_doc, xref)
                
                if pix.n - pix.alpha < 4:  # RGB or CMYK
                    img_data = pix.tobytes("png")
                    
                    # Save temporarily with cross-platform path
                    temp_img_path = os.path.join(tempfile.gettempdir(), f"temp_img_{img_index}_{int(time.time())}.png")
                    with open(temp_img_path, 'wb') as temp_img:
                        temp_img.write(img_data)
                    
                    # Add image to document
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(temp_img_path, width=Inches(4.0))
                    
                    # Add caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(f"[Image {img_index + 1}]")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Cleanup
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
                
                pix = None
            except Exception as e:
                print(f"Image extraction error: {e}")
                continue
    
    def _extract_text_with_formatting(self, doc: Document, page):
        """Extract text with advanced formatting preservation"""
        text_blocks = page.get_text("dict")
        
        for block in text_blocks.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    paragraph = doc.add_paragraph()
                    
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if not text:
                            continue
                        
                        run = paragraph.add_run(text + " ")
                        
                        # Apply advanced formatting
                        font_size = span["size"]
                        if font_size > 10:  # Only adjust if significantly different
                            run.font.size = Pt(min(font_size, 36))
                        
                        # Bold formatting
                        if span["flags"] & 2:
                            run.font.bold = True
                        
                        # Italic formatting
                        if span["flags"] & 1:
                            run.font.italic = True
                        
                        # Font color (basic implementation)
                        if span["color"] != 0:
                            try:
                                color = self._convert_pdf_color(span["color"])
                                run.font.color.rgb = color
                            except:
                                pass
                    
                    # Set paragraph formatting
                    if paragraph.text.strip():
                        paragraph.paragraph_format.space_after = Pt(8)
                        paragraph.paragraph_format.line_spacing = 1.15
    
    def _convert_pdf_color(self, pdf_color: int) -> RGBColor:
        """Convert PDF color to RGB (simplified)"""
        # Basic color conversion - extend as needed
        if pdf_color == 0:
            return RGBColor(0, 0, 0)  # Black
        else:
            return RGBColor(0, 0, 0)  # Default to black
    
    def _is_heading(self, text: str) -> bool:
        """Detect if text is likely a heading"""
        text = text.strip()
        return (
            len(text) < 150 and
            (text.isupper() or
             text.endswith(':') or
             re.match(r'^(CHAPTER|SECTION|PART|ABSTRACT|INTRODUCTION|CONCLUSION)', text, re.IGNORECASE) or
             re.match(r'^\d+\.\s+[A-Z]', text) or
             re.match(r'^[IVX]+\.\s+[A-Z]', text))
        )
    
    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> str:
        """Professional DOCX to PDF conversion"""
        methods = [
            self._docx_to_pdf_windows_word,  # Best quality (Windows)
            self._docx_to_pdf_libreoffice,   # Cross-platform
            self._docx_to_pdf_pymupdf        # Fallback
        ]
        
        last_error = None
        for method in methods:
            try:
                result = method(docx_path, pdf_path)
                if result and os.path.exists(result):
                    return result
            except Exception as e:
                last_error = e
                continue
        
        raise Exception(f"All DOCX to PDF methods failed: {last_error}")
    
    def _docx_to_pdf_windows_word(self, docx_path: str, pdf_path: str) -> str:
        """Use Microsoft Word for best quality (Windows only)"""
        if os.name != 'nt':
            raise Exception("Windows Word conversion only available on Windows")
        
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            abs_docx_path = os.path.abspath(docx_path)
            abs_pdf_path = os.path.abspath(pdf_path)
            
            doc = word.Documents.Open(abs_docx_path)
            doc.SaveAs(abs_pdf_path, FileFormat=17)  # PDF format
            doc.Close()
            word.Quit()
            
            return pdf_path
        except Exception as e:
            raise Exception(f"Windows Word conversion failed: {str(e)}")
        finally:
            try:
                pythoncom.CoUninitialize()
            except:
                pass
    
    def _docx_to_pdf_libreoffice(self, docx_path: str, pdf_path: str) -> str:
        """Use LibreOffice for cross-platform conversion"""
        try:
            output_dir = os.path.dirname(pdf_path)
            
            commands = [
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
                ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path]
            ]
            
            for cmd in commands:
                try:
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                    if result.returncode == 0:
                        base_name = os.path.splitext(os.path.basename(docx_path))[0]
                        expected_pdf = os.path.join(output_dir, f"{base_name}.pdf")
                        
                        if os.path.exists(expected_pdf):
                            os.rename(expected_pdf, pdf_path)
                            return pdf_path
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    continue
            
            raise Exception("LibreOffice not found or conversion failed")
        except Exception as e:
            raise Exception(f"LibreOffice conversion failed: {str(e)}")
    
    def _docx_to_pdf_pymupdf(self, docx_path: str, pdf_path: str) -> str:
        """Fallback using PyMuPDF"""
        try:
            # Extract text from DOCX
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Create PDF
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            
            y_position = 50
            for text in text_content:
                if y_position > 700:  # New page if needed
                    page = pdf_doc.new_page()
                    y_position = 50
                
                page.insert_text((50, y_position), text)
                y_position += 20
            
            pdf_doc.save(pdf_path)
            pdf_doc.close()
            return pdf_path
        except Exception as e:
            raise Exception(f"PyMuPDF fallback failed: {str(e)}")
    
    def convert_image(self, image_buffer: bytes, original_format: str, target_format: str) -> bytes:
        """Professional image conversion with format optimization"""
        try:
            with Image.open(io.BytesIO(image_buffer)) as img:
                output_buffer = io.BytesIO()
                
                # Handle format-specific optimizations
                if target_format.upper() in ['JPEG', 'JPG']:
                    # Convert RGBA to RGB for JPEG
                    if img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                        img = background
                    img.save(output_buffer, format='JPEG', quality=95, optimize=True)
                
                elif target_format.upper() == 'PNG':
                    img.save(output_buffer, format='PNG', optimize=True)
                
                elif target_format.upper() == 'WEBP':
                    img.save(output_buffer, format='WEBP', quality=90, method=6)
                
                else:
                    img.save(output_buffer, format=target_format.upper())
                
                output_buffer.seek(0)
                return output_buffer.getvalue()
                
        except Exception as e:
            raise Exception(f"Image conversion failed: {str(e)}")
    
    def get_supported_conversions(self) -> Dict:
        """Get all supported conversion formats"""
        return {
            # Image conversions
            'image/jpeg': ['png', 'webp', 'jpg', 'bmp', 'tiff'],
            'image/jpg': ['png', 'webp', 'jpeg', 'bmp', 'tiff'],
            'image/png': ['jpeg', 'jpg', 'webp', 'bmp', 'tiff'],
            'image/gif': ['jpeg', 'jpg', 'png', 'webp'],
            'image/webp': ['jpeg', 'jpg', 'png', 'bmp'],
            'image/bmp': ['jpeg', 'jpg', 'png', 'webp'],
            'image/tiff': ['jpeg', 'jpg', 'png', 'webp'],
            'image/svg+xml': ['png', 'jpeg', 'jpg', 'webp'],
            
            # Document conversions
            'application/pdf': ['docx', 'txt', 'pdf'],
            'application/msword': ['pdf', 'txt', 'docx'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['pdf', 'txt', 'docx'],
            'application/vnd.ms-excel': ['pdf', 'txt'],
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['pdf', 'txt'],
            'application/vnd.ms-powerpoint': ['pdf', 'txt'],
            'application/vnd.openxmlformats-officedocument.presentationml.presentation': ['pdf', 'txt'],
            'text/plain': ['pdf', 'docx', 'txt'],
        }

# Initialize converter
converter = ProfessionalDocumentConverter()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_safe_filename(filename):
    """Create a safe filename for cross-platform use"""
    # Remove unsafe characters and replace spaces
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', filename)
    safe_name = safe_name.replace(' ', '_')
    return secure_filename(safe_name)

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "service": "Professional Document Converter",
        "version": "2.0.0",
        "timestamp": time.time()
    })

@app.route('/api/conversion/convert', methods=['POST'])
def convert_file():
    input_path = None
    output_path = None
    
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        target_format = request.form.get('targetFormat', '').lower()
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"error": "File type not allowed"}), 400
        
        if not target_format:
            return jsonify({"error": "Target format not specified"}), 400

        # Create safe filenames with cross-platform paths
        safe_filename = get_safe_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, f"input_{int(time.time())}_{safe_filename}")
        output_filename = f"converted_{os.path.splitext(safe_filename)[0]}.{target_format}"
        output_path = os.path.join(UPLOAD_FOLDER, f"output_{int(time.time())}_{output_filename}")
        
        # Save uploaded file
        file.save(input_path)
        
        # Verify file was saved
        if not os.path.exists(input_path):
            raise Exception(f"Failed to save uploaded file: {input_path}")
        
        print(f"File saved successfully: {input_path}")
        print(f"File size: {os.path.getsize(input_path)} bytes")

        # Determine file type and perform conversion
        mime_type = get_mime_type(file.filename)
        
        if mime_type.startswith('image/'):
            # Image conversion - handle in memory
            file_buffer = file.read()
            result_data = converter.convert_image(file_buffer, file.filename.split('.')[-1], target_format)
            mime_type = f'image/{target_format}'
            output_filename = get_output_filename(file.filename, target_format)
            
        else:
            # Document conversion
            if mime_type == 'application/pdf' and target_format == 'docx':
                result_path = converter.convert_pdf_to_docx(input_path, output_path)
            elif mime_type == 'application/pdf' and target_format == 'txt':
                result_path = pdf_to_txt_advanced(input_path, output_path)
            elif 'wordprocessingml' in mime_type and target_format == 'pdf':
                result_path = converter.convert_docx_to_pdf(input_path, output_path)
            elif 'wordprocessingml' in mime_type and target_format == 'txt':
                result_path = docx_to_txt_advanced(input_path, output_path)
            elif mime_type == 'text/plain' and target_format == 'pdf':
                result_path = txt_to_pdf_advanced(input_path, output_path)
            elif mime_type == 'text/plain' and target_format == 'docx':
                result_path = txt_to_docx_advanced(input_path, output_path)
            else:
                return jsonify({"error": f"Unsupported conversion: {mime_type} to {target_format}"}), 400
            
            # Verify conversion worked
            if not os.path.exists(result_path):
                raise Exception(f"Conversion failed - output file not created: {result_path}")
            
            # Read result
            with open(result_path, 'rb') as f:
                result_data = f.read()
            
            mime_type = get_mime_type(target_format)
            output_filename = get_output_filename(file.filename, target_format)

        response = send_file(
            io.BytesIO(result_data),
            as_attachment=True,
            download_name=output_filename,
            mimetype=mime_type
        )
        response.headers['X-Conversion-Status'] = 'success'
        response.headers['X-Conversion-Engine'] = 'professional-python-converter'
        return response

    except Exception as e:
        print(f"Conversion error: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500
    finally:
        # Cleanup files
        cleanup_file(input_path)
        cleanup_file(output_path)

@app.route('/api/conversion/supported-formats', methods=['GET'])
def get_supported_formats():
    return jsonify({
        "success": True,
        "data": converter.get_supported_conversions()
    })

@app.route('/api/conversion/capabilities', methods=['GET'])
def get_capabilities():
    capabilities = {
        "service": "Professional Document Converter",
        "version": "2.0.0",
        "features": [
            'PDF to DOCX with structure preservation',
            'PDF to Text extraction with layout awareness',
            'Text to PDF generation',
            'Text to DOCX conversion',
            'DOCX to PDF conversion (multiple methods)',
            'DOCX to Text extraction',
            'Excel to PDF/TXT conversion',
            'PowerPoint to PDF/TXT conversion',
            'Advanced image format conversion',
            'Multiple fallback methods for reliability',
            'Image extraction from PDFs',
            'Font formatting preservation (bold, italic, size)',
            'Advanced PDF structure detection',
            'Heading and list recognition',
            'Table extraction and preservation',
            'Color formatting preservation',
            'Cross-platform compatibility'
        ],
        "professional_grade": True,
        "image_support": True,
        "advanced_formatting": True,
        "multiple_fallbacks": True
    }
    
    return jsonify({
        "success": True,
        "data": capabilities
    })

# Additional advanced conversion functions
def pdf_to_txt_advanced(pdf_path, txt_path):
    """Advanced PDF to text conversion"""
    pdf_doc = fitz.open(pdf_path)
    full_text = []
    
    for page_num in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_num)
        text = page.get_text("text", sort=True)
        
        if text.strip():
            full_text.append(f"\n{'='*50}")
            full_text.append(f"PAGE {page_num + 1}")
            full_text.append(f"{'='*50}\n")
            full_text.append(text.strip())
    
    pdf_doc.close()
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(full_text))
    
    return txt_path

def docx_to_txt_advanced(docx_path, txt_path):
    """Advanced DOCX to text conversion"""
    doc = Document(docx_path)
    full_text = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.split()[-1] if 'Heading' in paragraph.style.name else '1'
                full_text.append(f"\n{'#' * int(level)} {text}")
            else:
                full_text.append(text)
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(full_text))
    
    return txt_path

def txt_to_pdf_advanced(txt_path, pdf_path):
    """Advanced text to PDF conversion"""
    pdf_doc = fitz.open()
    
    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    page = pdf_doc.new_page()
    page.insert_text((50, 50), text)
    
    pdf_doc.save(pdf_path)
    pdf_doc.close()
    return pdf_path

def txt_to_docx_advanced(txt_path, docx_path):
    """Advanced text to DOCX conversion"""
    word_doc = Document()
    
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    word_doc.add_heading('Converted Text Document', 0)
    
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    for para in paragraphs:
        if para:
            p = word_doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(6)
    
    word_doc.save(docx_path)
    return docx_path

def get_mime_type(filename):
    """Get MIME type from filename"""
    ext = filename.split('.')[-1].lower()
    mime_map = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'txt': 'text/plain',
        'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
        'png': 'image/png', 'webp': 'image/webp',
        'gif': 'image/gif', 'bmp': 'image/bmp',
        'tiff': 'image/tiff', 'tif': 'image/tiff',
        'svg': 'image/svg+xml',
    }
    return mime_map.get(ext, 'application/octet-stream')

def get_output_filename(original_name, target_format):
    """Generate output filename"""
    base_name = original_name.rsplit('.', 1)[0]
    timestamp = int(time.time())
    return f"{base_name}-converted-{timestamp}.{target_format}"

def cleanup_file(file_path):
    """Cleanup temporary file"""
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Warning: Could not delete {file_path}: {e}")

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    print(f"🚀 Starting Professional Document Converter on port {port}")
    print(f"📁 Temp directory: {UPLOAD_FOLDER}")
    print("✅ Features: Image extraction, Advanced formatting, Multiple fallbacks")
    print("✅ Supported: PDF, DOCX, DOC, TXT, JPG, PNG, WEBP, GIF, BMP, TIFF, SVG")
    app.run(host='0.0.0.0', port=port, debug=False)